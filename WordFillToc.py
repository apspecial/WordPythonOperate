import docx


def get_titles_and_paragraph_numbers(file_path):
    doc = docx.Document(file_path)
    titles_and_paragraph_numbers = []
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name.startswith('Heading'):
            titles_and_paragraph_numbers.append((paragraph.text.strip(), index + 1))
    return titles_and_paragraph_numbers


def display_content_between_paragraphs(file_path, start_paragraph, end_paragraph):
    doc = docx.Document(file_path)
    content = []

    for i in range(start_paragraph - 1, end_paragraph):
        content.append(doc.paragraphs[i].text)

    return ''.join(content)


def find_element(lst, element):
    for index, value in enumerate(lst):
        if value == element:
            return index
    return -1


def insert_paragraph(doc, text, position):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    paragraph.insert_paragraph_before(paragraph, position)


def copy_paragraphs(src_doc, dest_doc, src_start, src_end, dest_index):
    src_paragraphs = src_doc.paragraphs[src_start:src_end]
    dest_paragraphs = dest_doc.paragraphs[dest_index:dest_index + 1]

    for i, src_paragraph in enumerate(src_paragraphs):
        dest_paragraph = dest_paragraphs[0].insert_paragraph_before()
        dest_paragraph.add_run(src_paragraph.text)


def scan_file_fill_toc(in_file, toc_file, out_file):
    source_doc = docx.Document(in_file)
    # 计算段落总数
    paragraph_count = len(source_doc.paragraphs)
    # 查找第一个文档的标题
    titles_and_paragraph_numbers = get_titles_and_paragraph_numbers(in_file)
    # 所有的标题
    all_title = [title for title, paragraph_number in titles_and_paragraph_numbers]
    last_index = len(all_title) - 1
    doc = docx.Document(toc_file)
    i = 0
    wr_ind = 0  # 写入多少段落，目标文档需要进行偏移
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            # 去掉空格
            para_text = paragraph.text.strip()
            if para_text in all_title:
                start_ind = find_element(all_title, para_text)  # 查找相同标题的段落
                tmp_ind = titles_and_paragraph_numbers[start_ind]
                start_para = tmp_ind[1]
                if start_ind < last_index:
                    tmp_ind = titles_and_paragraph_numbers[start_ind + 1]
                    end_para = tmp_ind[1]
                else:
                    end_para = paragraph_count + 1
                # 获取源文档
                paragraphs_to_copy = source_doc.paragraphs[start_para:end_para - 1]
                dest_paragraphs = doc.paragraphs[i + 1 + wr_ind:i + 2 + wr_ind]
                # 将段落添加到目标文档中
                for paragrapha in paragraphs_to_copy:
                    dest_paragraph = dest_paragraphs[0].insert_paragraph_before()
                    dest_paragraph.add_run(paragrapha.text)
                    wr_ind += 1
        i += 1
    # 保存目标文档
    doc.save(out_file)


if __name__ == "__main__":
    in_file = 'test.docx'  # 源文档
    toc_file = 'toc.docx'  # 目录文档
    out_file = 'toc1.docx'  # 拷贝后的文档
    scan_file_fill_toc(in_file, toc_file, out_file)
