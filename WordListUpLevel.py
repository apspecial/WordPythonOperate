# 将多级列表所有的标题级别提升M个级别
# 将多级列表所有的标题级别下降M个级别
import docx


def multi_list_up_level(in_file, out_file, m):  # M =1 代表所有级别向上提1级。
    doc_out = docx.Document(in_file)  # 输出文档， 以input_file为模版
    for k in range(len(doc_out.paragraphs) - 1):
        paragraph = doc_out.paragraphs[k]
        for i in range(9):
            title_text = str(i + 1) + '级标题'  # 生成 1级标题，2级标题，。。。，9级标题
            if paragraph.style.name.startswith(title_text):
                # doc_out.add_heading(paragraph.text, level=i)
                style = str(i + 1 - m) + '级标题'
                paragraph.style = style
                print(paragraph.text)
    doc_out.save(out_file)
    print("转换成功！")


def multi_list_down_level(in_file, out_file, m):  # M =1 代表所有级别向上提1级。
    doc_out = docx.Document(in_file)  # 输出文档， 以input_file为模版
    for k in range(len(doc_out.paragraphs) - 1):
        paragraph = doc_out.paragraphs[k]
        for i in range(9, 0, -1):  # 逆序，否者的话，每次循环都能检测到，每次都修改
            title_text = str(i + 1) + '级标题'  # 生成 1级标题，2级标题，。。。，9级标题
            if paragraph.style.name.startswith(title_text):
                style = str(i + 1 + m) + '级标题'
                if i + m + 1 > 9:
                    style = "标书正文"

                paragraph.style = style
                print(paragraph.text)
    doc_out.save(out_file)
    print("转换成功！")


def extract_custom_styled_titles(docx_file):
    doc = docx.Document(docx_file)
    custom_styled_titles = []
    for paragraph in doc.paragraphs:
        for i in range(9):
            title_text = str(i + 1) + '级标题'  # 生成 1级标题，2级标题，。。。，9级标题
            if paragraph.style.name.startswith(title_text):
                custom_styled_titles.append(title_text + ':' + paragraph.text)
    return custom_styled_titles


if __name__ == '__main__':
    # 设置参数
    input_file = 'temp_empty.docx'
    output_file = 'output1.docx'

    # multi_list_up_level(input_file, output_file, 1)
    multi_list_down_level(input_file, output_file, 3)
    # print(extract_custom_styled_titles(input_file))
