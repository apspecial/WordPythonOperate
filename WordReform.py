import docx
import re
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_custom_styled_titles(docx_file):
    doc = docx.Document(docx_file)
    custom_styled_titles = []
    for paragraph in doc.paragraphs:
        for i in range(9):
            title_text = str(i + 1) + '级标题'  # 生成 1级标题，2级标题，。。。，9级标题
            if paragraph.style.name.startswith(title_text):
                custom_styled_titles.append(title_text + ':' + paragraph.text)
    return custom_styled_titles


def read_word_file(file_path):  # 从一个word文档中读出所有内容。
    doca = docx.Document(file_path)
    full_text = []
    for para in doca.paragraphs:
        full_text.append(para.text)
    return ''.join(full_text)


def find_pattern_ind(input_text, pattern):  # 在一段字符串中，查找pattern所匹配的正则表达式出现的位置
    result2 = re.finditer(pattern, input_text)
    list_ind = []
    for match in result2:
        end = match.end()
        list_ind.append(end)
    return list_ind

    # 第二个列表按照第一个列表中元素的位置进行分割


def split_list(list1, list2):
    positions = [list2.index(i) for i in list1 if i in list2]
    lists = [list2[positions[i]:positions[i + 1]] for i in range(len(positions) - 1)]
    lists.append(list2[positions[-1]:])
    return positions, lists


# 使用count()方法, 计算零的个数
def count_zeros_before_index(lst, index):
    sub_lst = lst[:index]
    return sub_lst.count(0)


def remove_last_char(s):
    return s[:-1]


def find_title_reform(input_file, output_file):  # input_file为从输入文件,output_file是将最终结果输出的文档名称

    doc_in = docx.Document(input_file)  # 打开文档
    custom_styled_titles = []
    doc_out = docx.Document(input_file)  # 输出文档， 以input_file为模版
    doc_out._body.clear_content()
    doc_out.add_heading("技术指标符合性", level=3)
    pattern1 = r'\(\d+\)'  # 查找(1),(2)
    pattern2 = r'\d+\)'  # 查找1)
    # for paragraph in doc_in.paragraphs:        #查找文本段落内容, 每段文本
    for i in range(len(doc_in.paragraphs) - 1):
        paragraph = doc_in.paragraphs[i]
        if paragraph.style.name.startswith('4级标题'):
            doc_out.add_heading(paragraph.text, level=4)
        elif paragraph.style.name.startswith('标书正文'):  # 非标题, 正文内容

            match1 = re.findall(pattern1, paragraph.text)

            next_paragraph = doc_in.paragraphs[i + 1]

            # print(next_paragraph)
            match_next = re.findall(pattern1, next_paragraph.text)  # 下一段也找到 (2)

            if match1 and (not match_next):  # match1非空, 下一段为空
                write_text = paragraph.text.replace(''.join(match1), '')  # 将查找到的（1） 删除
                doc_out.add_heading(write_text, level=5)
                # 添加段落内容
                paragraph_add = doc_out.add_paragraph()
                paragraph_add.style = '标书正文'
                # # 设置段落的对齐方式
                # paragraph_add.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                # 添加文本并加粗
                run = paragraph_add.add_run("指标要求：")
                run.bold = True

                # # 对下一个段落进行操作
                # next_paragraph = paragraph._element.getnext()
                # tmp_match = re.findall(pattern2, next_paragraph)
                # if next_paragraph
            elif match1 and match_next:  # 下一段为非空
                write_text = remove_last_char(paragraph.text.replace(''.join(match1), '')) + '要求'  # 将查找到的（1） 删除
                doc_out.add_heading(write_text, level=5)
                # 添加段落内容
                paragraph_add = doc_out.add_paragraph()
                paragraph_add.style = '标书正文'
                # # 设置段落的对齐方式
                # paragraph_add.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                # 添加文本并加粗
                run = paragraph_add.add_run(
                    "指标要求：" + remove_last_char(paragraph.text.replace(''.join(match1), '')) + "。")
                run.bold = True

            else:
                match2 = re.findall(pattern2, paragraph.text)
                if match2:  # match1为空， match2非空
                    # 添加段落内容
                    paragraph_add = doc_out.add_paragraph()
                    paragraph_add.style = '标书正文'
                    # 添加文本并加粗
                    run = paragraph_add.add_run('(' + paragraph.text)
                    run.bold = True
    doc_out.save(output_file)
    print("转换成功！")


if __name__ == '__main__':
    # 设置参数
    input_file = 'temp_empty.docx'
    output_file = 'output.docx'

    find_title_reform(input_file, output_file)
