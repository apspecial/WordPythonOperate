from docx import Document
from docx.oxml.shared import qn
from docx.shared import Pt


# doc = Document()
# para = doc.add_paragraph('', style='List Number 3')
# para.paragraph_format.space_after = Pt(0)
# run = para.add_run('表内容1')
# run.font.name = '宋体'
# run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# run.font.bold = True
# run.font.size = Pt(14)


from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt


doc = Document('input.docx')
para = doc.add_paragraph()
pPr = para._element.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle', {qn('w:val'), 'lv3'}) # lv3 是从styles.xml里找到的对应样式w:styleId的值
pPr.append(pStyle)
para.add_run('表内容1')

# 保存文档
doc.save('demo.docx')