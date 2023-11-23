from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# 创建一个新的文档对象
doc = Document()

# 添加一个新的段落
paragraph = doc.add_paragraph()

# 设置段落的对齐方式为左对齐
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 在段落中添加一个项目符号列表
bullet_list = paragraph.add_run()

# 添加列表项
for i in range(5):
    bullet_list.add_bullet_point()
    run = bullet_list.add_run('Item ' + str(i+1))
    run.font.size = Pt(14)

# 保存文档
doc.save('demo.docx')
