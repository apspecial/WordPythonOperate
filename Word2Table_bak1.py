import docx
import re
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 将多级列表的参数的指标编号变成表格样式。


# 设置参数
input_file = 'input.docx'
# table_width = "\\textwidth"


# Create an instance of a word document
doc = docx.Document()


# def read_table(docx_file):
#     # 读取docx文件
#     doc = docx.Document(docx_file)


def read_word_file(file_path):
    doca = docx.Document(file_path)
    full_text = []
    for para in doca.paragraphs:
        full_text.append(para.text)
    return ''.join(full_text)


content = read_word_file(input_file)
#print(content[1:8])

# 查找一级标题，类似(1),(2)... 注意括号不能是中文格式
# text = "这里有一些数字：1(2)3，45(6)，7(8)9"
pattern = r'\(\d+\)'
result = re.findall(pattern, content)
#print(result)
# 需要存储一级标题的结束位置
one_class_s_end = []
for match in result:
    start = content.index(match)
    one_class_s_end.append(start)
    end = start + len(match)
    one_class_s_end.append(end)
    #print(f"匹配子串：{match}，起始位置：{start}，结束位置：{end}")
print(one_class_s_end)

# 查找二级标题1) 2).....， 这种方式查找会把一级标题也找到，因为一级标题包含1）的结构
pattern = r'\d+\)'

# 计算所有条目内容所需要的行数
row_num = len(re.findall(pattern, content)) - int(len(one_class_s_end) / 2) + 1
#print(row_num)

# Creating a table object, 技术指标复合型一览表。
table = doc.add_table(rows=row_num, cols=5)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = '序号'
row[1].text = '技术指标参数要求'
row[4].text = '是否符合'
# 合并第一行的单元格
table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3))

# # 合并第二行的后两列单元格
# table.cell(1, 1).merge(table.cell(1, 2))

# 循环查找匹配项
# 存储循环查找的上一个匹配的结束位置
last_end = 0
result2 = re.finditer(pattern, content)
#print(result2)
# row_iter = 1  #从第2行开始
two_class_end = []
for match in result2:
    match_txt = match.group()
    start = match.start()
    end = match.end()
    two_class_end.append(end)

print(two_class_end)
# 第二行
row_iter = 1
# 记录需要合并的行数
merge_row_num = 0
for item in two_class_end:
    row_it = table.rows[row_iter].cells
    pos_item = two_class_end.index(item)
    # #print(pos_item)
    if item != two_class_end[-1]:
        if item in one_class_s_end:  # 处理一级标题
            cont_ind_start = item
            cont_ind_end = two_class_end[pos_item + 1] - 2
            # 是否是一级标题，如果是，在第三列填入内容。行号不增加
            row_it[2].text = content[cont_ind_start:cont_ind_end]
        else:  # 是否是一级标题，如果不是，在第四列填入内容。行号增加
            cont_ind_start = item
            if two_class_end[pos_item + 1] in one_class_s_end:  # 判断下一个元素是不是一级标题，如果是取内容位置发生变化，而且需要合并单元格
                cont_ind_end = two_class_end[pos_item + 1] - 3
                # table.merge_cells(start_row=row_iter-merge_row_num+1, start_column=2, end_row=row_iter+1, end_column=2)
                print(merge_row_num)
                for i in range(merge_row_num+1):
                    table.cell(row_iter - merge_row_num + 1 + i, 2).merge(
                        table.cell(row_iter - merge_row_num + 2 + i, 2))
                merge_row_num = 0
            else:
                cont_ind_end = two_class_end[pos_item + 1] - 2
                merge_row_num += 1  # 需要合并的行数加1
            row_it[3].text = content[cont_ind_start:cont_ind_end]
            row_it[4].text = "符合"
            row_iter += 1
    else:
        row_it[3].text = content[item:]
        row_it[4].text = "符合"

# 设置表格属性

# 设置表格的样式，让所有单元格都有边框
table.style = 'Table Grid'
table.style.font.size = Pt(15)  # 字体大小15磅
# table.style.font.color.rgb = RGBColor.from_string("6495ED")  # 字体颜色
table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

# # 遍历表格的每一个单元格
# for row in table.rows:
#     for cell in row.cells:
#         # 设置单元格的文字居中
#         paragraph = cell.paragraphs[0]
#         # run = paragraph.runs[0]
#         # run.font.size = Pt(14)
#         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# 设置表格的字体大小,对齐方式
# for row in table.rows:
#     for cell in row.cells:
#         for paragraph in cell.paragraphs:
#             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             for run in paragraph.runs:
#                 run.font.size = Pt(12)

# Add a Title to the document
# doc.add_heading('GeeksForGeeks', 0)


# Now save the document to a location
doc.save('gfg.docx')
