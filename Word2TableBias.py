import docx
import re
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 将多级列表的参数的指标编号变成表格样式。


# 设置参数
input_file = 'input.docx'
# table_width = "\\textwidth"


# Create an instance of a word document
doc = docx.Document()


def read_word_file(file_path):  # 从一个word文档中读出所有内容。
    doca = docx.Document(file_path)
    full_text = []
    for para in doca.paragraphs:
        full_text.append(para.text)
    return ''.join(full_text)


def find_pattern_ind(input_text, pattern):  # 在一段字符串中，查找pattern所匹配的正则表达式出现的位置
    result2 = re.finditer(pattern, input_text)
    list_ind = []
    for match in result2:
        end = match.end()
        list_ind.append(end)
    return list_ind


content = read_word_file(input_file)
# print(content)

# 查找一级标题，类似(1),(2)... 注意括号不能是中文格式
# text = "这里有一些数字：1(2)3，45(6)，7(8)9"
pattern = r'\(\d+\)'
one_class_s_end = find_pattern_ind(content, pattern)

# 查找二级标题1) 2).....， 这种方式查找会把一级标题也找到，因为一级标题包含1）的结构
pattern = r'\d+\)'

# 计算所有条目内容所需要的行数
# row_num = len(re.findall(pattern, content)) - int(len(one_class_s_end) / 2) + 1
# row_num = len(re.findall(pattern, content)) - len(one_class_s_end) + 1
row_num = len(re.findall(pattern, content))  # 稍后计算准确行数
# print(row_num)

# Creating a table object, 技术指标复合型一览表。
table = doc.add_table(rows=row_num, cols=8)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = '序号'
row[1].text = '项目'
row[3].text = '技术指标参数要求'
row[4].text = '技术指标参数响应'
row[5].text = '偏离情况'
row[6].text = '详细说明正/父偏离情况'
row[7].text = '可查询的具体位置'
# 合并第一行的单元格
table.cell(0, 1).merge(table.cell(0, 2))

# 循环查找匹配项
# 存储循环查找的上一个匹配的结束位置
last_end = 0
result2 = re.finditer(pattern, content)
two_class_end = []
for match in result2:
    match_txt = match.group()
    start = match.start()
    end = match.end()
    two_class_end.append(end)


# 此函数去掉字符串最后的冒号,包括中英文格式的
def remove_colon_if_exists(s):
    if s.endswith(':'):
        return s[:-1]
    elif s.endswith('：'):
        return s[:-1]
    else:
        return s


# print(two_class_end)
# 第二行
row_iter = 1

# 序号
ind = 1
for item in two_class_end:
    # #print(item, row_iter)
    row_it = table.rows[row_iter].cells
    pos_item = two_class_end.index(item)
    # ##print(pos_item)
    if item != two_class_end[-1]:
        if item in one_class_s_end:  # 处理一级标题
            cont_ind_start = item
            cont_ind_end = two_class_end[pos_item + 1] - 2
            # 是否是一级标题，如果是，在第三列填入内容。如果下一个标题不是一级标题，行号不增加,否则行号增加
            if two_class_end[two_class_end.index(item) + 1] in one_class_s_end:  # 下一个是一级标题
                fill_text = remove_colon_if_exists(content[cont_ind_start:cont_ind_end - 1].strip())  # 删除前后的空格
                table.cell(row_iter, 1).merge(table.cell(row_iter, 2))  # 合并
                # 判断一级标题所填的内容是否有冒号，如果有冒号，分两列填写，否者合并这两列，填写一个cell
                if "：" in fill_text:
                    spli_txt = fill_text.split("：")  # 按冒号分割
                    row_it[1].text = spli_txt[0]
                    row_it[3].text = spli_txt[1]
                else:
                    row_it[3].text = fill_text
                row_it[4].text = "我所承诺" + row_it[3].text
                row_it[0].text = str(ind)
                row_it[5].text = "无偏离"
                row_iter += 1
                ind += 1
            else:
                row_it[1].text = remove_colon_if_exists(content[cont_ind_start:cont_ind_end].strip())  # 去掉冒号和空格
                # print(row_it[1].text)


        else:  # 是否是一级标题，如果不是，在第四列填入内容。行号增加
            cont_ind_start = item
            if two_class_end[pos_item + 1] in one_class_s_end:  # 判断下一个元素是不是一级标题，如果是,取内容位置发生变化，而且需要合并单元格
                cont_ind_end = two_class_end[pos_item + 1] - 3
            else:
                cont_ind_end = two_class_end[pos_item + 1] - 2
            fill_text = content[cont_ind_start:cont_ind_end].strip()
            if "：" in fill_text:
                spli_txt = fill_text.split("：")  # 按冒号分割
                row_it[2].text = spli_txt[0]
                row_it[3].text = spli_txt[1]
            else:
                row_it[3].text = fill_text
                # 合并两列
                table.cell(row_iter, 1).merge(table.cell(row_iter, 2))
            row_it[4].text = "我所承诺" + row_it[3].text
            row_it[0].text = str(ind)
            row_it[5].text = "无偏离"
            row_iter += 1
            ind += 1
    else:  # 处理最后一行
        row_it[3].text = content[item:].strip()
        row_it[4].text = "我所承诺" + row_it[3].text
        row_it[5].text = "无偏离"
        row_it[0].text = str(ind)
        table.cell(row_iter, 1).merge(table.cell(row_iter, 2))

if len(one_class_s_end) == 1: #只有一个一级标题的合并
    start_pos = 1
    for i in range(len(two_class_end) - 2):
        table.cell(i + 1, 1).merge(table.cell(i + 2, 1))
else:
    for i in range(len(one_class_s_end)):
        start_pos = two_class_end.index(one_class_s_end[i])
        merge_iter_num = two_class_end

for item in one_class_s_end:
    one_index = two_class_end.index(item)


# 第二个列表按照第一个列表中元素的位置进行分割
def split_list(list1, list2):
    positions = [list2.index(i) for i in list1 if i in list2]
    lists = [list2[positions[i]:positions[i + 1]] for i in range(len(positions) - 1)]
    lists.append(list2[positions[-1]:])
    return positions, lists


multi_list_tu = split_list(one_class_s_end, two_class_end)
pos_multi = multi_list_tu[0]
# print(pos_multi)
pos_multi_shift = pos_multi[1:]
pos_multi_shift.append(len(two_class_end))
# merge_len = pos_multi_shift - pos_multi
# 计算出需要合并的行数
merge_len = [a - b - 1 for a, b in zip(pos_multi_shift, pos_multi)]

merge_start = [a + 1 for a in pos_multi]


# 使用count()方法, 计算零的个数
def count_zeros_before_index(lst, index):
    sub_lst = lst[:index]
    return sub_lst.count(0)


# print(merge_start)
# print(merge_len)
need_sub = 0
for i in range(len(merge_len)):
    merge_start[i] -= need_sub
    if merge_len[i] != 0:
        need_sub += 1

# print(merge_start)

for i in range(len(merge_start)):  # 需要合并几个大块
    # print(f"i={i}")
    for j in range(merge_len[i] - 1):  # 每一块要合并多少次
        # print(f"j={j}, merge_start[i] = {merge_start[i]}")
        table.cell(merge_start[i] + j, 1).merge(table.cell(merge_start[i] + j + 1, 1))

# 设置表格属性

# 设置表格的样式，让所有单元格都有边框
table.style = 'Table Grid'
table.style.font.size = Pt(14)  # 字体大小15磅
# table.style.font.color.rgb = RGBColor.from_string("6495ED")  # 字体颜色
table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 对齐

# 从后向前遍历表格的每一行，删除最后的空行
for row in reversed(table.rows):
    # 如果一行是空的
    if all(cell.text.strip() == '' for cell in row.cells):  #迭代查询所有的cell是否为空
        # 删除这一行
        row._element.getparent().remove(row._element)

# Now save the document to a location
# 此次是转换的技术条款偏离表
doc.save('gfg_bias.docx')
print("转换成功！")
