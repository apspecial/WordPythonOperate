from docx import Document

doc = Document("temp_empty.docx")

doc._body.clear_content()

doc.add_heading("段落1", level=3)

doc.add_heading("段落2", level=4)

doc.add_heading("段落3", level=4)

doc.add_heading("段落4", level=3)

doc.add_heading("段落5", level=4)

doc.add_heading("段落6", level=3)

doc.save("result.docx")